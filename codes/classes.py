import requests
import pandas as pd
from collections import Counter
from docx import Document
from docx.shared import Inches
from matplotlib import pyplot as plt

class MyApi():
    def __init__(self):
        self.url = "https://api.openalex.org/"
        pass
    
    
    def show_works(self, start_year, end_year=None, check_interrupt=None):
        """Méthode pour regrouper les publications attribuables à une institution donnée sur une période donnée"""
        url = self.__url_works_generator(start_year, end_year)
        results = self.__extract_data(url, check_interrupt)
        self.__generate_excel_file(results)
        return
    
    def show_collaborators(self, start_year, end_year=None, check_interrupt=None):
        """Méthode pour programmer l’extraction de la liste des pays collaborateurs pour la période."""
        country_counts = self.__extract_collaborators(start_year, end_year, check_interrupt)
        self.__collaborators_excel_file(country_counts)
        pass
    
    
    def generate_country_report(self, start_year, end_year=None, check_interrupt=None):
        """Méthode pour générer le document word contenant le graphique représentant les 10 principaux pays collaborateurs"""
        country_counts = self.__extract_collaborators(start_year, end_year, check_interrupt)
        self.__generate_graph(country_counts)
        self.__insert_into_word(start_year, end_year)
        
        
    def show_works_with_collaboration(self, collaborator_ror, start_year, end_year=None, check_interrupt=None):
        """Méthode pour programmer l’extraction de la liste des principaux sujets des publications en collaboration entre l’ÉTS et le CNRS"""
        target_rors = {
            "https://ror.org/0020snb74",
            f"{collaborator_ror}"
        }
        
        url = self.__url_works_generator(start_year, end_year)
        results = self.__extract_data(url, check_interrupt)
        
        # Premier filtrage pour obtenir uniquement les données associées aux ROR qui nous intéressent
        new_data = [
            work for work in results 
            if target_rors.issubset(
                {institution.get("ror") 
                for authorship in work.get("authorships", []) 
                for institution in authorship.get("institutions", [])
                if institution.get("ror")}
            )
        ]
        
        topics_list = []

        # Extraire les pays des institutions des co-auteurs
        for pub in new_data:
            for topic in pub.get("topics", []):
                topics_list.append(topic.get("display_name"))

        # Compter le nombre d'apparitions de chaque pays
        topic_counts = Counter(topics_list)
        """Ensuite, faisons l'extraction sous forme d'un graphe"""
        try:
            self.__generate_graph_topics(topic_counts)
        except ValueError as e:
            raise ValueError(e)
        return
    

        

# ******************************************************************************************************************
# ******************************************************************************************************************
# ******************************************************************************************************************
# ************************************** METHODES PRIVEES***********************************************************
# ******************************************************************************************************************
# ******************************************************************************************************************
# ******************************************************************************************************************


    def __url_works_generator(self, start_year, end_year=None):
        """Génération de l'url qui sera utilisée pour l'appel d'api"""
        publication_year = self.__generate_publication_year_filter(start_year, end_year)
        url= self.url + f'works?filter=authorships.institutions.ror:0020snb74,publication_year:{publication_year}&select=display_name,publication_year,doi,authorships,topics'
        return url
    
    def __extract_data(self, url, check_interrupt=None):
        """Extraction des données à partir de l'url générée"""
        all_results = []
        cursor = "*"  # Premier curseur pour la pagination

        while cursor and not (check_interrupt and check_interrupt()):
            paginated_url = f"{url}&per-page=200&cursor={cursor}"
            response = requests.get(paginated_url)

            if response.status_code != 200:
                print("❌ Erreur lors de la récupération des données:", response.status_code)
                break

            data = response.json()
            publications = data.get("results", [])

            meta = data.get("meta", {})

            if not publications:
                break  # Fin de la pagination

            all_results.extend(publications)
            cursor = meta.get("next_cursor")  # Mettre à jour le curseur

            # Vérification pour ne pas boucler infiniment
            if cursor is None:
                break
            
            # Vérification en cas d'interruption de l'opération
            if check_interrupt:
                check_interrupt()

        # Nombre total de publications récupérées
        print(f"Nombre total de publications récupérées : {len(all_results)}")           
        return all_results
        
    def __extract_collaborators(self, start_year, end_year=None, check_interrupt=None):
        """Méthode privée pour l’extraction de la liste des pays collaborateurs pour la période."""
        url = self.__url_works_generator(start_year, end_year)
        publications = self.__extract_data(url, check_interrupt)
        country_list = []

        # Extraire les pays des institutions des co-auteurs
        for pub in publications:
            for authorship in pub.get("authorships", []):
                for institution in authorship.get("institutions", []):
                    country = institution.get("country_code")
                    if country:  
                        country_list.append(country)  # Ajouter le pays s'il est présent

        # Compter le nombre d'apparitions de chaque pays
        country_counts = Counter(country_list)

        return country_counts
    
    def __collaborators_excel_file(self, country_counts):
         # Convertir en DataFrame pour l'export Excel
        df = pd.DataFrame(country_counts.items(), columns=["Pays", "Nombre de publications"])
        df.sort_values(by="Nombre de publications", ascending=False, inplace=True)

        # Sauvegarder le fichier Excel
        df.to_excel("pays_collaborateurs_ets.xlsx", index=False, engine="openpyxl")
        print("✅ La liste des pays collaborateurs a été enregistrée dans 'pays_collaborateurs_ets.xlsx'.")
    
    def __generate_publication_year_filter(self, start_year, end_year=None):
        """Génère la chaîne de filtre publication_year pour l'API OpenAlex."""
        if end_year is None:  # Si une seule année est donnée
            return f"{start_year}"
        
        if start_year > end_year:
            raise ValueError("L'année de début doit être inférieure ou égale à l'année de fin.")

        years = "|".join(str(year) for year in range(start_year, end_year + 1))
 
        return f"{years}"
    
    def __generate_graph(self, country_counts):
        # Obtenir les 10 principaux pays
        top_countries = country_counts.most_common(10)

        countries, counts = zip(*top_countries)

        # Créer le graphique
        plt.figure(figsize=(10, 5))
        plt.bar(countries, counts, color='royalblue')
        plt.xlabel("Pays")
        plt.ylabel("Nombre de publications")
        plt.title("Top 10 des pays collaborateurs")
        plt.xticks(rotation=45)
        plt.grid(axis='y', linestyle="--", alpha=0.7)

        # Sauvegarder le graphique
        plt.savefig("top_countries.png", bbox_inches="tight")
        plt.close()
        
        
    def __generate_graph_topics(self, topics_count):
        # Obtenir les 10 principaux pays
        top_topics = topics_count.most_common(20)

        try:
            topics, counts = zip(*top_topics)

            # Créer le graphique
            plt.figure(figsize=(12, 6))
            plt.bar(topics, counts, color='royalblue')
            plt.xlabel("Sujets")
            plt.ylabel("Nombre de publications")
            plt.title("Top 20 des sujets de collaboration avec l'ÉTS")
            plt.xticks(rotation=30, ha="right", fontsize=10)
            plt.grid(axis='y', linestyle="--", alpha=0.7)

            # Sauvegarder le graphique
            plt.savefig("top_topics.png", bbox_inches="tight")
            plt.close()
            print("✅ Graphique généré avec succès: top_topics.png")
        
        except(ValueError):
            raise ValueError("ROR INVALIDE")
        

    def __insert_into_word(self, start_year, end_year=None):
        """Génère un rapport word avec le graphe généré"""
        doc = Document()
        doc.add_heading("Analyse des collaborations internationales", level=1)

        doc.add_paragraph(f"Le graphique ci-dessous présente les 10 principaux pays ayant collaboré avec l'ÉTS sur les publications de {start_year} à {end_year}.")

        doc.add_picture("top_countries.png", width=Inches(6))

        doc.save("rapport_collaborations.docx")
        print("✅ Rapport Word généré avec succès : rapport_collaborations.docx")

            
    def __generate_excel_file(self, all_results):
        
        """Génère un fichier excel à partir des résultats obtenus"""
        # Extraire les données dans une liste de dictionnaires
        publications_data = [
            {
                "Titre": pub.get("display_name", "Non disponible"),
                "Année": pub.get("publication_year", "Non disponible"),
                "Lien vers l'article": pub.get("doi", "Non disponible"),
            }
            for pub in all_results
        ]

        # Convertir en DataFrame
        df = pd.DataFrame(publications_data)

        # Trier les publications par année (ordre croissant)
        df['Année'] = pd.to_numeric(df['Année'], errors='coerce')  
        df.sort_values(by='Année', ascending=True, inplace=True)

        # Sauvegarder sous Excel avec plusieurs feuilles (sheets)
        filename = "publications.xlsx"
        batch_size = 1000  # Nombre de publications par feuille

        with pd.ExcelWriter(filename, engine='openpyxl') as writer:
            for i in range(0, len(df), batch_size):
                sheet_name = f"Sheet_{(i // batch_size) + 1}"
                df.iloc[i:i + batch_size].to_excel(writer, sheet_name=sheet_name, index=False)

        print(f"✅ Les données ont été enregistrées dans '{filename}', divisées en plusieurs feuilles.")

            
